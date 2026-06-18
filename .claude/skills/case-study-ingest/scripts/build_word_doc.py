#!/usr/bin/env python
"""
Build a conformant Whitewater case-study Word document (Output 1) from a
case-study JSON entry, by cloning the locked template so all styling/branding
is preserved.

Template (source of truth):
    assets/cs-word/Case Study Template.docx

The generator is deliberately "dumb": it renders exactly what is in the JSON
entry. All provenance markers ([DRAFT - for review], [AI - verify],
[TO COMPLETE]) must already be present in the JSON field values. `date_reviewed`
is left to the reviewing RP and is rendered blank unless explicitly set.

Usage:
    python build_word_doc.py --json case-studies.json --id "<entry id>" --out "assets/cs-word/"
    python build_word_doc.py --json case-studies.json --index 0 --out out.docx

Requires: python-docx
"""
import argparse
import copy
import json
import os
import re
import sys

import docx

# --- Field map: template placeholder paragraph text -> (json key, kind) -------
# "kind": narrative = single paragraph body; meta = "Label: value" line.
TEMPLATE_PATH = os.path.join("assets", "cs-word", "Case Study Template.docx")

# Section body placeholders in the template (exact text), mapped to JSON keys.
BODY_PLACEHOLDERS = {
    "Single line headline": "hook",
    "One paragraph narrative summary of ambition": "ambition",
    "One paragraph narrative summary of the challenge": "challenge",
    "One paragraph narrative of solutions applied.": "solutions",
    "One paragraph narrative of client insights.": "client_insight",
}

# Client Details lines: label prefix (before ':') -> JSON key.
DETAIL_FIELDS = [
    ("Company name", "client_name"),
    ("Industry", "industry"),
    ("Company size, revenue", "revenue"),
    ("Company size, employees", "headcount"),
    ("Geography", "geography"),
    ("Type of reinvention", "type_of_reinvention"),
    ("Duration", "duration"),
    ("Client lead for reinvention (e.g., CEO, CFO, CHRO)", "client_lead"),
    ("Do we have client permission to use their name", "_permission"),
    ("Whitewater point of contact", "ww_contact"),
    ("Date reviewed", "date_reviewed"),
]

RESULTS_PLACEHOLDER = "Bullet point results."
TITLE_PLACEHOLDER = "<Name>"


def set_paragraph_text(p, text):
    """Replace a paragraph's text while preserving its first run's formatting."""
    runs = p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(text)


def permission_value(entry):
    """Render the permission line. website_approved True -> Yes, else No (default)."""
    if entry.get("website_approved") is True:
        return "Yes"
    return "No"


def fill_narrative_field(doc, placeholder, value):
    """Find the placeholder paragraph and expand to support multi-paragraph values.
    Values containing \\n\\n are split and each part cloned from the anchor paragraph
    to inherit its formatting (style, font, spacing)."""
    for p in doc.paragraphs:
        if p.text.strip() == placeholder:
            anchor = p
            break
    else:
        return

    parts = [s.strip() for s in value.split('\n\n') if s.strip()] if value else []
    if not parts:
        set_paragraph_text(anchor, value or "")
        return

    set_paragraph_text(anchor, parts[0])
    prev = anchor
    for part in parts[1:]:
        new_p = copy.deepcopy(anchor._p)
        prev._p.addnext(new_p)
        new_para = docx.text.paragraph.Paragraph(new_p, anchor._parent)
        set_paragraph_text(new_para, part)
        prev = new_para


def _set_bullet_spacing(p):
    """Set 6 pt space-after on a bullet paragraph and disable contextualSpacing.
    The List Paragraph style has contextualSpacing enabled, which suppresses
    space_after between consecutive same-style paragraphs — must be explicitly
    overridden with val='0' at the paragraph level."""
    from docx.shared import Pt
    from lxml import etree
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.find(f'{{{W}}}pPr')
    if pPr is None:
        pPr = etree.SubElement(p._p, f'{{{W}}}pPr')
    ctx = pPr.find(f'{{{W}}}contextualSpacing')
    if ctx is None:
        ctx = etree.SubElement(pPr, f'{{{W}}}contextualSpacing')
    ctx.set(f'{{{W}}}val', '0')


def fill_results(doc, benefits):
    """Find the single bullet placeholder and expand to one bullet per benefit."""
    for p in doc.paragraphs:
        if p.text.strip() == RESULTS_PLACEHOLDER:
            anchor = p
            break
    else:
        raise ValueError("Results bullet placeholder not found in template")

    if not benefits:
        set_paragraph_text(anchor, "[TO COMPLETE] results to be supplied by RP")
        return

    set_paragraph_text(anchor, benefits[0])
    # Set spacing + disable contextualSpacing BEFORE deepcopy so all clones inherit both
    _set_bullet_spacing(anchor)

    prev = anchor
    for b in benefits[1:]:
        new_p = copy.deepcopy(anchor._p)
        prev._p.addnext(new_p)
        new_para = docx.text.paragraph.Paragraph(new_p, anchor._parent)
        set_paragraph_text(new_para, b)
        prev = new_para


def fill_details(doc, entry):
    """Replace each Client Details line: label part bold, value part not bold."""
    from docx.oxml.ns import qn
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def _clear_bold(r_elem):
        """Remove <w:b> and <w:bCs> from a run element's rPr."""
        rpr = r_elem.find(f'{{{W}}}rPr')
        if rpr is None:
            return
        for tag in ('b', 'bCs'):
            el = rpr.find(f'{{{W}}}{tag}')
            if el is not None:
                rpr.remove(el)

    by_prefix = {prefix: key for prefix, key in DETAIL_FIELDS}
    for p in doc.paragraphs:
        txt = p.text.strip()
        if ":" not in txt:
            continue
        prefix = txt.split(":", 1)[0].strip()
        if prefix not in by_prefix:
            continue
        key = by_prefix[prefix]
        if key == "_permission":
            value = permission_value(entry)
        else:
            value = entry.get(key, "") or ""

        runs = p.runs
        # Remove all but the first run; keep first to preserve font/size/colour.
        for r in runs[1:]:
            r._element.getparent().remove(r._element)

        if runs:
            # First run → label, bold
            runs[0].text = prefix
            runs[0].bold = True

            # Clone first run for the value (inherits font/size), then clear bold
            value_elem = copy.deepcopy(runs[0]._element)
            runs[0]._element.addnext(value_elem)
            # Set text in the cloned element
            t = value_elem.find(f'{{{W}}}t')
            if t is None:
                from lxml import etree
                t = etree.SubElement(value_elem, f'{{{W}}}t')
            value_text = f": {value}".rstrip()
            t.text = value_text
            if value_text and value_text[0] == ' ':
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            _clear_bold(value_elem)
        else:
            r1 = p.add_run(prefix)
            r1.bold = True
            r2 = p.add_run(f": {value}".rstrip())
            r2.bold = False

        _set_bullet_spacing(p)


def display_name(entry):
    """Best available name for the doc title and filename. Real client_name is
    preferred (Word is the internal source of truth), but when it is empty or an
    unfilled [TO COMPLETE] placeholder, fall back to public_client_name, then
    title. The placeholder still surfaces in the Client Details "Company name"
    line, so the gap stays flagged where it belongs."""
    cn = (entry.get("client_name") or "").strip()
    if not cn or cn.startswith("[TO COMPLETE]"):
        return (entry.get("public_client_name") or entry.get("title")
                or "case-study").strip()
    return cn


def build(entry, out_path):
    doc = docx.Document(TEMPLATE_PATH)

    # Title (<Name>) -> display name (real client_name, or public name when the
    # real one is an unfilled placeholder).
    for p in doc.paragraphs:
        if p.text.strip() == TITLE_PLACEHOLDER:
            set_paragraph_text(p, display_name(entry))
            break

    # Section bodies (multi-paragraph values split on \n\n).
    for placeholder, key in BODY_PLACEHOLDERS.items():
        fill_narrative_field(doc, placeholder, entry.get(key, "") or "")

    fill_results(doc, entry.get("benefits") or [])
    fill_details(doc, entry)

    doc.save(out_path)
    return out_path


def safe_name(s):
    """Filesystem-safe fragment: keep word chars, hyphen, en-dash, ampersand and
    spaces (all valid in Windows filenames); drop other punctuation and collapse
    whitespace so a removed character never leaves a double space."""
    s = re.sub(r"[^\w\-–& ]", "", s or "")
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def reinvention_label(entry):
    """Short reinvention label for filenames: the primary type only, cut at the
    first elaboration separator (comma or dash). Keeps auto-derived names bounded
    when type_of_reinvention holds a long descriptive string."""
    raw = (entry.get("type_of_reinvention") or "").strip()
    for sep in (",", " - ", " – "):
        idx = raw.find(sep)
        if idx != -1:
            raw = raw[:idx]
    return safe_name(raw) or "Reinvention"


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--json", required=True, help="Path to case-studies.json")
    ap.add_argument("--id", help="Entry id to build")
    ap.add_argument("--index", type=int, help="Entry index to build")
    ap.add_argument("--out", required=True, help="Output .docx path or directory")
    args = ap.parse_args()

    data = json.load(open(args.json, encoding="utf-8"))
    items = data if isinstance(data, list) else data.get("case_studies", data)

    if args.id is not None:
        entry = next((e for e in items if e.get("id") == args.id), None)
        if entry is None:
            sys.exit(f"No entry with id={args.id!r}")
    elif args.index is not None:
        entry = items[args.index]
    else:
        sys.exit("Provide --id or --index")

    out = args.out
    if os.path.isdir(out) or out.endswith(("/", "\\")):
        fname = f"Case Study {safe_name(display_name(entry))} – {reinvention_label(entry)}.docx"
        out = os.path.join(out, fname)

    build(entry, out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
